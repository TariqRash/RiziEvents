from __future__ import annotations

import hashlib
import re
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "RiziEvents-Technical-Documentation.docx"
TMP_DIR = ROOT / "tmp" / "docs" / "mermaid"


def ensure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    if "CodeBlock" not in document.styles:
        style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Courier New"
        style.font.size = Pt(8.5)
        pf = style.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.right_indent = Inches(0.25)
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)


def render_mermaid_png(source: str) -> Path:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    name = hashlib.sha256(source.encode("utf-8")).hexdigest()[:16]
    output_path = TMP_DIR / f"{name}.png"
    if output_path.exists():
        return output_path

    req = urllib.request.Request(
        "https://kroki.io/mermaid/png",
        data=source.encode("utf-8"),
        headers={
            "Content-Type": "text/plain",
            "Accept": "image/png",
            "User-Agent": "Mozilla/5.0",
        },
    )
    with urllib.request.urlopen(req, timeout=60) as response:
        output_path.write_bytes(response.read())
    return output_path


def add_markdown(document: Document, path: Path) -> None:
    in_code = False
    code_lang = ""
    mermaid_lines: list[str] = []

    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                mermaid_lines = []
                if code_lang and code_lang != "mermaid":
                    document.add_paragraph(code_lang.upper(), style="Intense Quote")
            else:
                if code_lang == "mermaid":
                    png_path = render_mermaid_png("\n".join(mermaid_lines).strip() + "\n")
                    document.add_picture(str(png_path), width=Inches(6.5))
                    document.add_paragraph("")
                in_code = False
                code_lang = ""
                mermaid_lines = []
            continue

        if in_code:
            if code_lang == "mermaid":
                mermaid_lines.append(line)
            else:
                document.add_paragraph(line or " ", style="CodeBlock")
            continue

        if not line.strip():
            document.add_paragraph("")
            continue

        iframe = re.match(r'^<iframe\b[^>]*src="([^"]+)"[^>]*></iframe>$', line)
        if iframe:
            document.add_paragraph(f"Embedded mockup reference: {iframe.group(1)}")
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        if re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue

        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue

        document.add_paragraph(line)


def main() -> None:
    document = Document()
    ensure_styles(document)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    document.add_heading("RiziEvents Technical Documentation", level=0)
    document.add_paragraph("Reduced MVP submission package for the Holberton phase-three review.")

    add_markdown(document, ROOT / "docs" / "technical-documentation.md")

    document.add_page_break()
    add_markdown(document, ROOT / "docs" / "api-spec.md")

    document.add_page_break()
    add_markdown(document, ROOT / "docs" / "scm-qa-plan.md")

    document.add_page_break()
    add_markdown(document, ROOT / "docs" / "technical-justifications.md")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
